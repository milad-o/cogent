"""Word document loader."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from agenticflow.document.loaders.base import BaseLoader
from agenticflow.document.types import Document


class WordLoader(BaseLoader):
    """Loader for Microsoft Word documents.
    
    Supports .docx files. Extracts paragraphs and tables.
    Note: .doc (old format) is not directly supported.
    
    Example:
        >>> loader = WordLoader()
        >>> docs = await loader.load(Path("document.docx"))
    """
    
    supported_extensions = [".docx", ".doc"]
    
    def __init__(
        self,
        encoding: str = "utf-8",
        include_tables: bool = True,
    ) -> None:
        """Initialize the loader.
        
        Args:
            encoding: Not used for Word docs (binary format).
            include_tables: Whether to extract table content.
        """
        super().__init__(encoding)
        self.include_tables = include_tables
    
    async def load(self, path: Path, **kwargs: Any) -> list[Document]:
        """Load a Word document.
        
        Args:
            path: Path to the Word document.
            **kwargs: Additional options.
            
        Returns:
            List containing a single Document.
            
        Raises:
            ImportError: If python-docx is not installed.
            ValueError: If file is .doc format (not supported).
        """
        if path.suffix.lower() == ".doc":
            raise ValueError(
                "Old .doc format is not supported. "
                "Please convert to .docx format."
            )
        
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "Word document loading requires 'python-docx'. "
                "Install with: uv add python-docx"
            )
        
        doc = DocxDocument(str(path))
        
        # Extract all paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)
        
        # Extract tables if enabled
        tables_text = []
        if self.include_tables:
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                tables_text.append("\n".join(rows))
            
            if tables_text:
                content += "\n\n" + "\n\n".join(tables_text)
        
        return [
            self._create_document(
                content,
                path,
                paragraph_count=len(paragraphs),
                table_count=len(doc.tables),
            )
        ]


__all__ = ["WordLoader"]
