"""
Document Loaders for AgenticFlow.

Native document loading supporting major file types:
- Text: .txt, .md, .rst
- Documents: .pdf, .docx, .doc
- Data: .csv, .json, .jsonl, .xlsx, .xls
- Web: .html, .htm
- Code: .py, .js, .ts, .java, .cpp, .c, .go, .rs, .rb, etc.

Example:
    >>> from agenticflow.retriever.loaders import DocumentLoader, load_documents
    >>> 
    >>> # Load a single file
    >>> loader = DocumentLoader()
    >>> docs = await loader.load("document.pdf")
    >>> 
    >>> # Load multiple files
    >>> docs = await loader.load_many(["doc1.pdf", "doc2.md", "data.csv"])
    >>> 
    >>> # Load a directory
    >>> docs = await loader.load_directory("./documents", glob="**/*.pdf")
    >>> 
    >>> # Quick function
    >>> docs = await load_documents("./data")
"""

from __future__ import annotations

import asyncio
import csv
import html
import io
import json
import mimetypes
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Sequence


@dataclass
class Document:
    """A loaded document with content and metadata.
    
    Attributes:
        content: The text content of the document.
        metadata: Metadata about the document (source, page, etc.).
    """
    
    content: str
    metadata: dict[str, Any] = field(default_factory=dict)
    
    def __post_init__(self) -> None:
        if "source" not in self.metadata:
            self.metadata["source"] = "unknown"
    
    def __repr__(self) -> str:
        preview = self.content[:50] + "..." if len(self.content) > 50 else self.content
        return f"Document(content='{preview}', source='{self.metadata.get('source', 'unknown')}')"


# =============================================================================
# File Type Handlers
# =============================================================================

async def _load_text(path: Path, encoding: str = "utf-8") -> list[Document]:
    """Load plain text files."""
    try:
        content = path.read_text(encoding=encoding)
    except UnicodeDecodeError:
        # Try with different encodings
        for enc in ["latin-1", "cp1252", "iso-8859-1"]:
            try:
                content = path.read_text(encoding=enc)
                break
            except UnicodeDecodeError:
                continue
        else:
            content = path.read_bytes().decode("utf-8", errors="replace")
    
    return [Document(
        content=content,
        metadata={
            "source": str(path),
            "filename": path.name,
            "file_type": path.suffix.lower(),
        }
    )]


async def _load_markdown(path: Path, encoding: str = "utf-8") -> list[Document]:
    """Load markdown files with optional frontmatter extraction."""
    content = path.read_text(encoding=encoding)
    metadata: dict[str, Any] = {
        "source": str(path),
        "filename": path.name,
        "file_type": ".md",
    }
    
    # Extract YAML frontmatter if present
    if content.startswith("---"):
        parts = content.split("---", 2)
        if len(parts) >= 3:
            frontmatter = parts[1].strip()
            content = parts[2].strip()
            # Parse simple key: value frontmatter
            for line in frontmatter.split("\n"):
                if ":" in line:
                    key, value = line.split(":", 1)
                    metadata[key.strip()] = value.strip()
    
    return [Document(text=content, metadata=metadata)]


async def _load_html(path: Path, encoding: str = "utf-8") -> list[Document]:
    """Load HTML files and extract text content."""
    raw = path.read_text(encoding=encoding)
    
    # Try BeautifulSoup first
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(raw, "html.parser")
        
        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()
        
        # Extract title
        title = soup.title.string if soup.title else None
        
        # Get text
        text = soup.get_text(separator="\n", strip=True)
    except ImportError:
        # Fallback to regex-based extraction
        title = None
        title_match = re.search(r"<title[^>]*>(.*?)</title>", raw, re.IGNORECASE | re.DOTALL)
        if title_match:
            title = html.unescape(title_match.group(1).strip())
        
        # Remove scripts and styles
        text = re.sub(r"<script[^>]*>.*?</script>", "", raw, flags=re.IGNORECASE | re.DOTALL)
        text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.IGNORECASE | re.DOTALL)
        # Remove all tags
        text = re.sub(r"<[^>]+>", " ", text)
        # Decode entities
        text = html.unescape(text)
        # Clean whitespace
        text = re.sub(r"\s+", " ", text).strip()
    
    metadata = {
        "source": str(path),
        "filename": path.name,
        "file_type": ".html",
    }
    if title:
        metadata["title"] = title
    
    return [Document(text=text, metadata=metadata)]


async def _load_pdf(path: Path) -> list[Document]:
    """Load PDF files using pypdf or pdfplumber."""
    documents: list[Document] = []
    
    # Try pypdf first (lighter weight)
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(str(path))
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                documents.append(Document(
                    content=text,
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "file_type": ".pdf",
                        "page": i + 1,
                        "total_pages": len(reader.pages),
                    }
                ))
        return documents
    except ImportError:
        pass
    
    # Try pdfplumber (better extraction)
    try:
        import pdfplumber
        
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    documents.append(Document(
                        content=text,
                        metadata={
                            "source": str(path),
                            "filename": path.name,
                            "file_type": ".pdf",
                            "page": i + 1,
                            "total_pages": len(pdf.pages),
                        }
                    ))
        return documents
    except ImportError:
        pass
    
    raise ImportError(
        "PDF loading requires 'pypdf' or 'pdfplumber'. "
        "Install with: pip install pypdf"
    )


async def _load_docx(path: Path) -> list[Document]:
    """Load DOCX files using python-docx."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError(
            "DOCX loading requires 'python-docx'. "
            "Install with: pip install python-docx"
        )
    
    doc = DocxDocument(str(path))
    
    # Extract all paragraphs
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = "\n\n".join(paragraphs)
    
    # Extract tables
    tables_text = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        tables_text.append("\n".join(rows))
    
    if tables_text:
        content += "\n\n" + "\n\n".join(tables_text)
    
    return [Document(
        content=content,
        metadata={
            "source": str(path),
            "filename": path.name,
            "file_type": ".docx",
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
        }
    )]


async def _load_csv(path: Path, encoding: str = "utf-8") -> list[Document]:
    """Load CSV files."""
    content = path.read_text(encoding=encoding)
    reader = csv.DictReader(io.StringIO(content))
    
    rows = list(reader)
    fieldnames = reader.fieldnames or []
    
    # Convert to readable text format
    text_rows = []
    for row in rows:
        row_text = ", ".join(f"{k}: {v}" for k, v in row.items() if v)
        text_rows.append(row_text)
    
    return [Document(
        content="\n".join(text_rows),
        metadata={
            "source": str(path),
            "filename": path.name,
            "file_type": ".csv",
            "columns": list(fieldnames),
            "row_count": len(rows),
        }
    )]


async def _load_json(path: Path, encoding: str = "utf-8") -> list[Document]:
    """Load JSON files."""
    content = path.read_text(encoding=encoding)
    data = json.loads(content)
    
    # Handle both single objects and arrays
    if isinstance(data, list):
        documents = []
        for i, item in enumerate(data):
            documents.append(Document(
                content=json.dumps(item, indent=2, ensure_ascii=False),
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "file_type": ".json",
                    "index": i,
                    "total_items": len(data),
                }
            ))
        return documents
    else:
        return [Document(
            content=json.dumps(data, indent=2, ensure_ascii=False),
            metadata={
                "source": str(path),
                "filename": path.name,
                "file_type": ".json",
            }
        )]


async def _load_jsonl(path: Path, encoding: str = "utf-8") -> list[Document]:
    """Load JSON Lines files."""
    content = path.read_text(encoding=encoding)
    lines = [line.strip() for line in content.split("\n") if line.strip()]
    
    documents = []
    for i, line in enumerate(lines):
        try:
            data = json.loads(line)
            documents.append(Document(
                content=json.dumps(data, indent=2, ensure_ascii=False),
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "file_type": ".jsonl",
                    "line": i + 1,
                    "total_lines": len(lines),
                }
            ))
        except json.JSONDecodeError:
            continue
    
    return documents


async def _load_xlsx(path: Path) -> list[Document]:
    """Load Excel files using openpyxl."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError(
            "Excel loading requires 'openpyxl'. "
            "Install with: pip install openpyxl"
        )
    
    wb = load_workbook(str(path), read_only=True, data_only=True)
    documents = []
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        rows = []
        headers: list[str] = []
        
        for i, row in enumerate(sheet.iter_rows(values_only=True)):
            if i == 0:
                headers = [str(c) if c else f"col_{j}" for j, c in enumerate(row)]
                continue
            
            if any(c is not None for c in row):
                row_text = ", ".join(
                    f"{headers[j]}: {c}" 
                    for j, c in enumerate(row) 
                    if c is not None and j < len(headers)
                )
                rows.append(row_text)
        
        if rows:
            documents.append(Document(
                content="\n".join(rows),
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "file_type": ".xlsx",
                    "sheet": sheet_name,
                    "columns": headers,
                    "row_count": len(rows),
                }
            ))
    
    wb.close()
    return documents


async def _load_code(path: Path, encoding: str = "utf-8") -> list[Document]:
    """Load source code files."""
    content = path.read_text(encoding=encoding)
    
    # Detect language from extension
    ext_to_lang = {
        ".py": "python",
        ".js": "javascript",
        ".ts": "typescript",
        ".jsx": "javascript",
        ".tsx": "typescript",
        ".java": "java",
        ".cpp": "cpp",
        ".c": "c",
        ".h": "c",
        ".hpp": "cpp",
        ".go": "go",
        ".rs": "rust",
        ".rb": "ruby",
        ".php": "php",
        ".swift": "swift",
        ".kt": "kotlin",
        ".scala": "scala",
        ".sql": "sql",
        ".sh": "shell",
        ".bash": "shell",
        ".zsh": "shell",
        ".yaml": "yaml",
        ".yml": "yaml",
        ".toml": "toml",
        ".xml": "xml",
    }
    
    ext = path.suffix.lower()
    language = ext_to_lang.get(ext, "text")
    
    return [Document(
        content=content,
        metadata={
            "source": str(path),
            "filename": path.name,
            "file_type": ext,
            "language": language,
            "line_count": content.count("\n") + 1,
        }
    )]


# =============================================================================
# File Type Registry
# =============================================================================

# Map extensions to loader functions
LOADERS: dict[str, Callable[..., Any]] = {
    # Text
    ".txt": _load_text,
    ".text": _load_text,
    ".rst": _load_text,
    ".log": _load_text,
    # Markdown
    ".md": _load_markdown,
    ".markdown": _load_markdown,
    # HTML
    ".html": _load_html,
    ".htm": _load_html,
    # PDF
    ".pdf": _load_pdf,
    # Office
    ".docx": _load_docx,
    ".doc": _load_docx,  # Will fail gracefully if .doc
    # Data
    ".csv": _load_csv,
    ".tsv": _load_csv,
    ".json": _load_json,
    ".jsonl": _load_jsonl,
    ".ndjson": _load_jsonl,
    ".xlsx": _load_xlsx,
    ".xls": _load_xlsx,  # Will fail if old format
    # Code
    ".py": _load_code,
    ".js": _load_code,
    ".ts": _load_code,
    ".jsx": _load_code,
    ".tsx": _load_code,
    ".java": _load_code,
    ".cpp": _load_code,
    ".c": _load_code,
    ".h": _load_code,
    ".hpp": _load_code,
    ".go": _load_code,
    ".rs": _load_code,
    ".rb": _load_code,
    ".php": _load_code,
    ".swift": _load_code,
    ".kt": _load_code,
    ".scala": _load_code,
    ".sql": _load_code,
    ".sh": _load_code,
    ".bash": _load_code,
    ".zsh": _load_code,
    ".yaml": _load_code,
    ".yml": _load_code,
    ".toml": _load_code,
    ".xml": _load_code,
    ".css": _load_code,
    ".scss": _load_code,
    ".less": _load_code,
}


# =============================================================================
# Document Loader
# =============================================================================

class DocumentLoader:
    """Universal document loader supporting multiple file types.
    
    Automatically detects file type and uses appropriate loader.
    Supports text, PDF, DOCX, CSV, JSON, Excel, HTML, and code files.
    
    Args:
        encoding: Default text encoding (default: utf-8).
        extra_loaders: Additional file type handlers.
        
    Example:
        >>> loader = DocumentLoader()
        >>> 
        >>> # Load single file
        >>> docs = await loader.load("report.pdf")
        >>> 
        >>> # Load multiple files
        >>> docs = await loader.load_many(["doc1.pdf", "data.csv"])
        >>> 
        >>> # Load directory with glob pattern
        >>> docs = await loader.load_directory("./docs", glob="**/*.md")
    """
    
    def __init__(
        self,
        encoding: str = "utf-8",
        extra_loaders: dict[str, Callable[..., Any]] | None = None,
    ):
        self.encoding = encoding
        self._loaders = {**LOADERS}
        if extra_loaders:
            self._loaders.update(extra_loaders)
    
    def register_loader(
        self,
        extension: str,
        loader: Callable[[Path], Any],
    ) -> None:
        """Register a custom loader for a file extension.
        
        Args:
            extension: File extension (e.g., ".xyz").
            loader: Async function that takes Path and returns list[Document].
        """
        if not extension.startswith("."):
            extension = f".{extension}"
        self._loaders[extension.lower()] = loader
    
    async def load(
        self,
        path: str | Path,
        **kwargs: Any,
    ) -> list[Document]:
        """Load a single file.
        
        Args:
            path: Path to file.
            **kwargs: Additional arguments passed to loader.
            
        Returns:
            List of documents (may be multiple for multi-page files).
            
        Raises:
            FileNotFoundError: If file doesn't exist.
            ValueError: If file type is not supported.
        """
        path = Path(path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        
        if not path.is_file():
            raise ValueError(f"Not a file: {path}")
        
        ext = path.suffix.lower()
        loader = self._loaders.get(ext)
        
        if not loader:
            # Try mime type detection
            mime_type, _ = mimetypes.guess_type(str(path))
            if mime_type and mime_type.startswith("text/"):
                loader = _load_text
            else:
                raise ValueError(
                    f"Unsupported file type: {ext}. "
                    f"Supported: {', '.join(sorted(self._loaders.keys()))}"
                )
        
        # Prepare kwargs
        loader_kwargs = {"encoding": self.encoding, **kwargs}
        
        # Remove encoding for binary loaders
        if ext in {".pdf", ".docx", ".doc", ".xlsx", ".xls"}:
            loader_kwargs.pop("encoding", None)
        
        # Call loader
        import inspect
        if inspect.iscoroutinefunction(loader):
            return await loader(path, **{
                k: v for k, v in loader_kwargs.items()
                if k in inspect.signature(loader).parameters
            })
        else:
            return loader(path, **{
                k: v for k, v in loader_kwargs.items()
                if k in inspect.signature(loader).parameters
            })
    
    async def load_many(
        self,
        paths: Sequence[str | Path],
        **kwargs: Any,
    ) -> list[Document]:
        """Load multiple files concurrently.
        
        Args:
            paths: List of file paths.
            **kwargs: Additional arguments passed to loaders.
            
        Returns:
            Combined list of all documents.
        """
        tasks = [self.load(p, **kwargs) for p in paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        documents = []
        for result in results:
            if isinstance(result, Exception):
                # Log or handle error
                continue
            documents.extend(result)
        
        return documents
    
    async def load_directory(
        self,
        directory: str | Path,
        glob: str = "**/*",
        recursive: bool = True,
        exclude: Sequence[str] | None = None,
        **kwargs: Any,
    ) -> list[Document]:
        """Load all files from a directory.
        
        Args:
            directory: Directory path.
            glob: Glob pattern for file matching (default: all files).
            recursive: Whether to search recursively.
            exclude: Patterns to exclude.
            **kwargs: Additional arguments passed to loaders.
            
        Returns:
            List of all loaded documents.
        """
        directory = Path(directory)
        
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        if not directory.is_dir():
            raise ValueError(f"Not a directory: {directory}")
        
        exclude = exclude or []
        exclude_patterns = [re.compile(p) for p in exclude]
        
        # Collect matching files
        files = []
        for path in directory.glob(glob):
            if not path.is_file():
                continue
            
            # Check exclusions
            rel_path = str(path.relative_to(directory))
            if any(p.search(rel_path) for p in exclude_patterns):
                continue
            
            # Check if we have a loader
            if path.suffix.lower() in self._loaders:
                files.append(path)
        
        return await self.load_many(files, **kwargs)


# =============================================================================
# Convenience Functions
# =============================================================================

async def load_documents(
    source: str | Path | Sequence[str | Path],
    **kwargs: Any,
) -> list[Document]:
    """Load documents from file(s) or directory.
    
    Convenience function that auto-detects source type.
    
    Args:
        source: File path, directory path, or list of paths.
        **kwargs: Additional arguments for loaders.
        
    Returns:
        List of loaded documents.
        
    Example:
        >>> # Load single file
        >>> docs = await load_documents("report.pdf")
        >>> 
        >>> # Load directory
        >>> docs = await load_documents("./documents")
        >>> 
        >>> # Load multiple files
        >>> docs = await load_documents(["doc1.pdf", "doc2.md"])
    """
    loader = DocumentLoader(**{k: v for k, v in kwargs.items() if k == "encoding"})
    other_kwargs = {k: v for k, v in kwargs.items() if k != "encoding"}
    
    if isinstance(source, (list, tuple)):
        return await loader.load_many(source, **other_kwargs)
    
    path = Path(source)
    if path.is_dir():
        return await loader.load_directory(path, **other_kwargs)
    else:
        return await loader.load(path, **other_kwargs)


def load_documents_sync(
    source: str | Path | Sequence[str | Path],
    **kwargs: Any,
) -> list[Document]:
    """Synchronous version of load_documents.
    
    Args:
        source: File path, directory path, or list of paths.
        **kwargs: Additional arguments for loaders.
        
    Returns:
        List of loaded documents.
    """
    return asyncio.run(load_documents(source, **kwargs))


__all__ = [
    "Document",
    "DocumentLoader",
    "load_documents",
    "load_documents_sync",
    # Individual loaders for customization
    "LOADERS",
]
